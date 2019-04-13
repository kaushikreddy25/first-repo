#! python3

import docx

d = docx.Document('W:\\Python\\Programs\\demo.docx')

print(d.paragraphs)

p0 = d.paragraphs[0]

print(p0.text)

p1 = d.paragraphs[1]

print(p1.text)

print(p1.runs[0].text)

print(p1.runs[1].text)
print(p1.runs[2].text)
print(p1.runs[3].text)

p1.runs[3].text = 'italics, bold and underlined.'
p1.runs[3].bold = True
p1.runs[3].italic = True
p1.runs[3].underline = True

p1.style = 'Title'

d.save('W:\\Python\\Programs\\demo2.docx')

d3= docx.Document()
d3.add_paragraph('This is new paragraph')
d3.add_paragraph('2nd Paragraph is here. Hello!')

p30 = d3.paragraphs[0]
p30.style='Title'

d3.save('W:\\Python\\Programs\\demo3.docx')

print('-------------------')
#TODO - Create a function to extract all text in a document to a string


def getText(fileName):
    doc = docx.Document(fileName)
    fullText = []
    
    for p in doc.paragraphs: 
        fullText.append(p.text)
    
    return '\n'.join(fullText)

print(getText('W:\\Python\\Programs\\Form.docx'))














